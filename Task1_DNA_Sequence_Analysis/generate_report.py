from docx import Document
from docx.shared import Inches, Pt, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.style import WD_STYLE_TYPE
import os

def create_report():
    doc = Document()

    # Page margins
    section = doc.sections[0]
    section.top_margin = Inches(1)
    section.bottom_margin = Inches(1)
    section.left_margin = Inches(1.2)
    section.right_margin = Inches(1.2)

    # Title Page
    title = doc.add_paragraph()
    title.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = title.add_run("DNA/PROTEIN SEQUENCE ANALYSIS")
    run.bold = True
    run.font.size = Pt(20)
    run.font.color.rgb = RGBColor(0, 102, 153)

    subtitle = doc.add_paragraph()
    subtitle.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run2 = subtitle.add_run("Human Insulin Gene — BLAST Analysis Report")
    run2.bold = True
    run2.font.size = Pt(14)

    doc.add_paragraph()
    info = doc.add_paragraph()
    info.alignment = WD_ALIGN_PARAGRAPH.CENTER
    info.add_run("Sharmi S\n").bold = True
    info.add_run("B.Sc. Artificial Intelligence & Machine Learning\n")
    info.add_run("Hindusthan College of Arts and Science, Coimbatore\n")
    info.add_run("CodeAlpha Bioinformatics Internship — Task 1\n")
    info.add_run("2026")

    doc.add_page_break()

    # Section 1: Introduction
    doc.add_heading("1. Introduction", level=1)
    doc.add_paragraph(
        "Bioinformatics is an interdisciplinary field that combines biology, computer science, "
        "and data analysis to understand biological data. One of the most fundamental tasks in "
        "bioinformatics is sequence analysis — the process of examining DNA or protein sequences "
        "to extract meaningful biological information."
    )
    doc.add_paragraph(
        "DNA (Deoxyribonucleic Acid) is the molecule that carries genetic information in all "
        "living organisms. It is composed of four nucleotide bases: Adenine (A), Thymine (T), "
        "Guanine (G), and Cytosine (C). The specific order of these bases encodes instructions "
        "for building proteins — the molecules that perform virtually every function in the body."
    )
    doc.add_paragraph(
        "Insulin is a hormone protein produced by the pancreas that regulates blood glucose levels. "
        "It is encoded by the INS gene located on chromosome 11. Studying the insulin gene sequence "
        "across species provides valuable insights into evolutionary conservation and functional "
        "importance — regions that remain identical across species are likely critical for function."
    )

    # Section 2: Objectives
    doc.add_heading("2. Objectives", level=1)
    objectives = [
        "Download the human insulin DNA sequence (NM_000207) from NCBI database",
        "Perform BLAST analysis to identify homologous sequences across species",
        "Document similarity percentages, identity scores, and E-values",
        "Compare human insulin directly with rat, mouse, pig, cow, and dog sequences",
        "Visualize evolutionary relationships through sequence identity charts",
        "Interpret results in the context of evolutionary biology and medical applications"
    ]
    for obj in objectives:
        p = doc.add_paragraph(obj, style="List Bullet")

    # Section 3: Methodology
    doc.add_heading("3. Methodology", level=1)

    doc.add_heading("3.1 Tools and Technologies", level=2)
    tools = [
        ("Python 3.14", "Core programming language for automation"),
        ("BioPython 1.87", "Biological sequence analysis library"),
        ("NCBI Entrez API", "Fetching sequences from NCBI database"),
        ("NCBI BLAST (blastn)", "Finding homologous sequences"),
        ("Matplotlib", "Data visualization and chart generation"),
        ("Git & GitHub", "Version control and project hosting")
    ]
    for tool, purpose in tools:
        p = doc.add_paragraph(style="List Bullet")
        run = p.add_run(f"{tool}: ")
        run.bold = True
        p.add_run(purpose)

    doc.add_heading("3.2 Data Collection", level=2)
    doc.add_paragraph(
        "The human insulin mRNA sequence was retrieved from the NCBI Nucleotide database using "
        "the accession ID NM_000207. This sequence represents the complete mRNA transcript of "
        "the human insulin gene, including both coding and non-coding regions. The sequence was "
        "downloaded in FASTA format and saved locally for analysis."
    )

    doc.add_heading("3.3 BLAST Analysis", level=2)
    doc.add_paragraph(
        "BLAST (Basic Local Alignment Search Tool) was used to search for homologous sequences "
        "in the NCBI nucleotide database (nt). The blastn program was selected for nucleotide-to-"
        "nucleotide comparison. The top 20 matches were retrieved and analyzed for identity "
        "percentage, E-value, and bit score."
    )

    doc.add_heading("3.4 Direct Species Comparison", level=2)
    doc.add_paragraph(
        "In addition to BLAST analysis, direct pairwise alignment was performed between human "
        "insulin and five other mammalian species: rat (NM_012582), mouse (NM_008386), pig "
        "(NM_214049), cow (NM_173932), and dog (NM_001159779). Global sequence alignment was "
        "performed using BioPython's pairwise2 module to calculate identity percentages."
    )

    # Section 4: Results
    doc.add_heading("4. Results", level=1)

    doc.add_heading("4.1 BLAST Analysis Results", level=2)
    doc.add_paragraph(
        "The BLAST search returned 20 significant matches. All matches showed E-values "
        "approaching zero, confirming that all identified sequences are genuine homologs "
        "rather than random matches. The top matches were predominantly primate species, "
        "reflecting the close evolutionary relationship between humans and other primates."
    )

    # BLAST results table
    table1 = doc.add_table(rows=10, cols=3)
    table1.style = "Table Grid"
    headers = ["Species", "Identity %", "E-value"]
    for i, header in enumerate(headers):
        cell = table1.rows[0].cells[i]
        cell.text = header
        cell.paragraphs[0].runs[0].bold = True

    blast_data = [
        ("Homo sapiens (insulin isoform)", "100.0%", "~0"),
        ("Homo sapiens (preproinsulin)", "99.56%", "~0"),
        ("Homo sapiens mRNA", "99.78%", "~0"),
        ("Gorilla gorilla", "98.28%", "~0"),
        ("Symphalangus syndactylus", "96.77%", "~0"),
        ("Pongo abelii (Orangutan)", "96.56%", "~0"),
        ("Hylobates moloch", "96.34%", "~0"),
        ("Nomascus leucogenys", "96.72%", "~0"),
        ("Macaca mulatta (Rhesus monkey)", "93.95%", "~0"),
    ]
    for i, (species, identity, evalue) in enumerate(blast_data):
        row = table1.rows[i + 1]
        row.cells[0].text = species
        row.cells[1].text = identity
        row.cells[2].text = evalue

    doc.add_paragraph()

    doc.add_heading("4.2 BLAST Visualization", level=2)
    doc.add_paragraph(
        "Figure 1 shows the sequence identity percentages for the top BLAST matches. "
        "All primate species show above 90% identity with human insulin, confirming "
        "strong evolutionary conservation of this gene."
    )
    doc.add_picture("results/identity_chart.png", width=Inches(5.5))
    doc.add_paragraph("Figure 1: BLAST Analysis — Sequence Identity % Across Primate Species").alignment = WD_ALIGN_PARAGRAPH.CENTER

    doc.add_heading("4.3 Direct Species Comparison Results", level=2)
    doc.add_paragraph(
        "Direct pairwise alignment between human insulin mRNA and five mammalian species "
        "revealed moderate sequence identity ranging from 59.8% to 64.0%. These values "
        "reflect full mRNA comparison including non-coding regions."
    )

    # Species comparison table
    table2 = doc.add_table(rows=6, cols=3)
    table2.style = "Table Grid"
    headers2 = ["Species", "Identity %", "Relationship"]
    for i, header in enumerate(headers2):
        cell = table2.rows[0].cells[i]
        cell.text = header
        cell.paragraphs[0].runs[0].bold = True

    species_data = [
        ("Cow", "64.0%", "Distant mammal"),
        ("Pig", "63.6%", "Distant mammal"),
        ("Rat", "63.2%", "Distant mammal"),
        ("Mouse", "61.2%", "Distant mammal"),
        ("Dog", "59.8%", "Distant mammal"),
    ]
    for i, (species, identity, rel) in enumerate(species_data):
        row = table2.rows[i + 1]
        row.cells[0].text = species
        row.cells[1].text = identity
        row.cells[2].text = rel

    doc.add_paragraph()
    doc.add_picture("results/species_comparison_chart.png", width=Inches(5.5))
    doc.add_paragraph("Figure 2: Direct Mammal Comparison — Identity % with Human Insulin").alignment = WD_ALIGN_PARAGRAPH.CENTER

    # Section 5: Discussion
    doc.add_heading("5. Discussion", level=1)
    doc.add_paragraph(
        "The BLAST analysis results clearly demonstrate the evolutionary conservation of the "
        "insulin gene across primate species. The extremely high identity percentages (93-100%) "
        "among primates reflect their recent common ancestry and the critical functional importance "
        "of insulin in glucose metabolism."
    )
    doc.add_paragraph(
        "The E-values approaching zero for all matches confirm that these are genuine biological "
        "homologs. An E-value represents the probability that a match occurred by random chance — "
        "values close to zero mean the similarity is almost certainly due to shared evolutionary "
        "ancestry rather than coincidence."
    )
    doc.add_paragraph(
        "The direct species comparison revealed lower identity percentages (59-64%) for non-primate "
        "mammals. This is expected because the comparison was performed on full mRNA sequences, "
        "which include untranslated regions (UTRs) that evolve much faster than protein-coding "
        "regions. At the protein level, pig insulin differs from human insulin by only one amino "
        "acid — which is why pig insulin was successfully used to treat human diabetes for decades "
        "before synthetic insulin was developed."
    )
    doc.add_paragraph(
        "This finding has significant medical implications. The high conservation of insulin "
        "across species demonstrates why animal-derived insulin can be used as a therapeutic "
        "substitute in humans. Understanding sequence conservation patterns is fundamental to "
        "drug development and cross-species medical research."
    )

    # Section 6: Conclusion
    doc.add_heading("6. Conclusion", level=1)
    doc.add_paragraph(
        "This project successfully demonstrated the use of bioinformatics tools to analyze "
        "DNA sequence similarity across species. The BLAST analysis of the human insulin gene "
        "revealed strong evolutionary conservation among primates and moderate conservation "
        "among other mammals. Key conclusions include:"
    )
    conclusions = [
        "Human insulin shows 93-100% sequence identity with primate species",
        "All BLAST matches had E-values approaching zero — confirming genuine homology",
        "Non-primate mammals show 59-64% mRNA identity due to non-coding region variation",
        "The high protein-level conservation explains the historical medical use of animal insulin",
        "Bioinformatics tools like BLAST are powerful for understanding evolutionary relationships"
    ]
    for c in conclusions:
        doc.add_paragraph(c, style="List Bullet")

    # Section 7: References
    doc.add_heading("7. References", level=1)
    references = [
        "NCBI Nucleotide Database. Human insulin mRNA (NM_000207). https://www.ncbi.nlm.nih.gov",
        "Altschul SF, et al. Basic local alignment search tool. J Mol Biol. 1990;215(3):403-10.",
        "Cock PJ, et al. Biopython: freely available Python tools for computational biology. Bioinformatics. 2009.",
        "UniProt Consortium. UniProt: the universal protein knowledgebase. Nucleic Acids Res. 2021.",
        "Hunter JD. Matplotlib: A 2D graphics environment. Computing in Science & Engineering. 2007."
    ]
    for ref in references:
        doc.add_paragraph(ref, style="List Bullet")

    # Save document
    output_path = "report/analysis_report.docx"
    doc.save(output_path)
    print(f"Report saved to {output_path}")
    print("Convert to PDF by opening in Word and Save As PDF!")

if __name__ == "__main__":
    create_report()
